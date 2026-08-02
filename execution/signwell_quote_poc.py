"""Create SignWell-ready solar quote documents.

The default is a local dry-run. Live creation is explicit and uses SignWell
test mode by default so POC documents do not count as real signature requests.
"""

from __future__ import annotations

import argparse
import base64
import json
import os
import uuid
from pathlib import Path
from typing import Any
from urllib import request
from urllib.error import HTTPError


try:
    from execution.pandadoc_quote_poc import build_quote_lines, write_quote_docx
except ModuleNotFoundError:
    from pandadoc_quote_poc import build_quote_lines, write_quote_docx


REPO_ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = REPO_ROOT / ".tmp" / "signwell_quote_poc"
DEFAULT_BASE_URL = "https://www.signwell.com/api/v1"
INK = "24313A"
MUTED = "64727A"
ORANGE = "E58B2A"
PALE = "F5F7F6"


def _shade(cell: Any, fill: str) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    properties = cell._tc.get_or_add_tcPr()
    shading = properties.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        properties.append(shading)
    shading.set(qn("w:fill"), fill)


def _cell_margins(cell: Any, top: int = 100, start: int = 140, bottom: int = 100, end: int = 140) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    properties = cell._tc.get_or_add_tcPr()
    margins = properties.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        properties.append(margins)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = margins.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def _set_cell_text(cell: Any, text: str, *, bold: bool = False, color: str = INK, size: int = 9, align: Any = None) -> None:
    from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt, RGBColor
    cell.text = ""
    paragraph = cell.paragraphs[0]
    if align is not None:
        paragraph.alignment = align
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "Aptos"
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    _cell_margins(cell)


def build_polished_quote_docx(path: Path, analysis: dict[str, Any], recipient_name: str) -> None:
    from docx import Document
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Cm, Pt, RGBColor
    fields = analysis.get("fields", {})
    estimate = analysis.get("estimate") or {}
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(1.35)
    section.bottom_margin = Cm(1.25)
    section.left_margin = Cm(1.65)
    section.right_margin = Cm(1.65)
    normal = doc.styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(9)
    normal.font.color.rgb = RGBColor.from_string(INK)

    header = section.header
    header_table = header.add_table(rows=1, cols=2, width=Cm(17.7))
    header_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    header_table.columns[0].width = Cm(10.7)
    header_table.columns[1].width = Cm(7)
    _set_cell_text(header_table.cell(0, 0), "FIRST CLIENT", bold=True, color=ORANGE, size=18)
    _set_cell_text(header_table.cell(0, 1), "FIRST CLIENT BV", color=MUTED, size=8, align=WD_ALIGN_PARAGRAPH.RIGHT)
    for cell in header_table.rows[0].cells:
        cell._tc.get_or_add_tcPr().append(OxmlElement("w:tcBorders"))

    title = doc.add_paragraph()
    title.paragraph_format.space_before = Pt(18)
    title.paragraph_format.space_after = Pt(2)
    run = title.add_run("VOORLOPIGE OFFERTE")
    run.bold = True
    run.font.name = "Aptos Display"
    run.font.size = Pt(25)
    run.font.color.rgb = RGBColor.from_string(INK)
    subtitle = doc.add_paragraph("Zonnepanelen en thuisbatterij")
    subtitle.paragraph_format.space_after = Pt(14)
    subtitle_run = subtitle.runs[0]
    subtitle_run.font.size = Pt(12)
    subtitle_run.font.color.rgb = RGBColor.from_string(ORANGE)

    info = doc.add_table(rows=1, cols=2)
    info.alignment = WD_TABLE_ALIGNMENT.CENTER
    info.autofit = False
    info.columns[0].width = Cm(9.2)
    info.columns[1].width = Cm(8.5)
    left = info.cell(0, 0)
    right = info.cell(0, 1)
    _shade(left, PALE)
    _shade(right, PALE)
    _set_cell_text(left, f"KLANT\n{recipient_name or fields.get('contact_name') or 'Onbekend'}\n{fields.get('installation_address') or 'Adres wordt nog bevestigd'}", bold=False, size=9)
    _set_cell_text(right, f"OFFERTEGEGEVENS\nNummer: POC-{uuid.uuid4().hex[:8].upper()}\nDatum: {analysis.get('received_at') or 'Vandaag'}\nGeldig: 30 dagen", size=9)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(15)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run("Uw voorstel in een oogopslag")
    r.bold = True
    r.font.size = Pt(13)
    r.font.color.rgb = RGBColor.from_string(INK)

    summary = doc.add_table(rows=2, cols=4)
    summary.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ["INSTALLATIE", "DAK / MONTAGE", "THUISBATTERIJ", "GEWENSTE TIMING"]
    values = [
        str(estimate.get("system_size_kw") or "Op basis van aanvraag") + " kWp",
        str(fields.get("roof_or_mount") or "Nog te bevestigen"),
        str(fields.get("battery_backup") or "Nog te bevestigen"),
        str(fields.get("timeline") or "In overleg"),
    ]
    for index, header in enumerate(headers):
        _shade(summary.cell(0, index), ORANGE)
        _set_cell_text(summary.cell(0, index), header, bold=True, color="FFFFFF", size=7, align=WD_ALIGN_PARAGRAPH.CENTER)
        _shade(summary.cell(1, index), PALE)
        _set_cell_text(summary.cell(1, index), values[index], size=8, align=WD_ALIGN_PARAGRAPH.CENTER)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(7)
    r = p.add_run("Raming van de investering")
    r.bold = True
    r.font.size = Pt(13)
    r.font.color.rgb = RGBColor.from_string(INK)
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, text in enumerate(("OMSCHRIJVING", "AANTAL", "BEDRAG")):
        _shade(table.cell(0, i), INK)
        _set_cell_text(table.cell(0, i), text, bold=True, color="FFFFFF", size=8, align=WD_ALIGN_PARAGRAPH.RIGHT if i == 2 else None)
    items = [
        ("Zonnepanelen en installatie", str(fields.get("panel_count") or "Volgens ontwerp"), estimate.get("solar_cost_usd")),
        ("Thuisbatterij", "Optioneel", estimate.get("battery_cost_usd")),
        ("Contingentie en afwerking", "1", estimate.get("contingency_usd")),
    ]
    for label, count, amount in items:
        cells = table.add_row().cells
        _set_cell_text(cells[0], label, size=9)
        _set_cell_text(cells[1], count, size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
        value = "EUR n.t.b." if amount is None else f"EUR {float(amount):,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")
        _set_cell_text(cells[2], value, size=9, align=WD_ALIGN_PARAGRAPH.RIGHT)
    total = estimate.get("estimated_total_usd")
    cells = table.add_row().cells
    for cell in cells:
        _shade(cell, PALE)
    _set_cell_text(cells[0], "INDICATIEF TOTAAL", bold=True, size=10)
    _set_cell_text(cells[1], "", size=10)
    total_text = "EUR n.t.b." if total is None else f"EUR {float(total):,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")
    _set_cell_text(cells[2], total_text, bold=True, color=ORANGE, size=11, align=WD_ALIGN_PARAGRAPH.RIGHT)

    note = doc.add_paragraph()
    note.paragraph_format.space_before = Pt(12)
    note.paragraph_format.space_after = Pt(4)
    nr = note.add_run("Belangrijk")
    nr.bold = True
    nr.font.color.rgb = RGBColor.from_string(ORANGE)
    body = doc.add_paragraph("Deze offerte is indicatief en wordt intern gecontroleerd. De definitieve prijs en technische configuratie worden bevestigd na een huisbezoek en controle van de situatie ter plaatse.")
    body.paragraph_format.space_after = Pt(10)
    body.runs[0].font.color.rgb = RGBColor.from_string(MUTED)

    sign = doc.add_table(rows=2, cols=2)
    sign.alignment = WD_TABLE_ALIGNMENT.CENTER
    _shade(sign.cell(0, 0), INK)
    _shade(sign.cell(0, 1), INK)
    _set_cell_text(sign.cell(0, 0), "VOOR AKKOORD", bold=True, color="FFFFFF", size=8)
    _set_cell_text(sign.cell(0, 1), "DATUM", bold=True, color="FFFFFF", size=8)
    _set_cell_text(sign.cell(1, 0), "{{signature:1:y:signature_1:150:36}}", color="FFFFFF", size=10)
    _set_cell_text(sign.cell(1, 1), "{{date:1:y:date_1:100:24}}", color="FFFFFF", size=10)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer.add_run("First Client BV")
    footer_run.font.name = "Aptos"
    footer_run.font.size = Pt(8)
    footer_run.font.color.rgb = RGBColor.from_string(MUTED)
    doc.save(path)


def load_dotenv(path: Path) -> None:
    if not path.exists():
        return
    for line in path.read_text(encoding="utf-8").splitlines():
        stripped = line.strip()
        if not stripped or stripped.startswith("#") or "=" not in stripped:
            continue
        key, value = stripped.split("=", 1)
        os.environ.setdefault(key.strip(), value.strip().strip('"').strip("'"))


def build_signwell_payload(
    document_name: str,
    docx_path: Path,
    recipient_email: str,
    recipient_name: str,
    *,
    draft: bool = True,
    test_mode: bool = True,
    embedded_signing: bool = True,
) -> dict[str, Any]:
    return {
        "name": document_name,
        "subject": "Uw offerte voor zonnepanelen",
        "message": "Beste,\n\nU kunt uw offerte online bekijken en ondertekenen via de link in deze e-mail.",
        "files": [{"name": docx_path.name, "file_base64": base64.b64encode(docx_path.read_bytes()).decode("ascii")}],
        "recipients": [{"id": "1", "name": recipient_name, "email": recipient_email}],
        "draft": draft,
        "test_mode": test_mode,
        "embedded_signing": embedded_signing,
        "text_tags": True,
        "language": "nl",
        "metadata": {"source": "Orelix Office Offerte Assistent", "quote_id": uuid.uuid4().hex[:12]},
    }


def api_request(method: str, url: str, api_key: str, body: bytes | None = None) -> dict[str, Any]:
    headers = {"X-Api-Key": api_key, "Accept": "application/json", "Content-Type": "application/json"}
    req = request.Request(url, data=body, headers=headers, method=method)
    try:
        with request.urlopen(req, timeout=60) as response:
            raw = response.read().decode("utf-8")
            return json.loads(raw) if raw else {}
    except HTTPError as exc:
        detail = exc.read().decode("utf-8", errors="replace")
        raise RuntimeError(f"SignWell API error {exc.code}: {detail}") from exc


def create_document(api_key: str, base_url: str, payload: dict[str, Any]) -> dict[str, Any]:
    body = json.dumps(payload).encode("utf-8")
    return api_request("POST", f"{base_url}/documents", api_key, body)


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="SignWell solar quote signing POC")
    parser.add_argument("--analysis-file", type=Path, required=True)
    parser.add_argument("--recipient-email", required=True)
    parser.add_argument("--recipient-name", required=True)
    parser.add_argument("--out-dir", type=Path, default=OUT_DIR)
    parser.add_argument("--live-create", action="store_true")
    parser.add_argument("--send", action="store_true", help="Create a sent signing request instead of a draft")
    parser.add_argument("--test-mode", action="store_true", help="Explicitly use SignWell test mode")
    parser.add_argument("--embedded", action="store_true", help="Use embedded signing instead of SignWell email delivery")
    return parser.parse_args()


def main() -> None:
    args = parse_args()
    load_dotenv(REPO_ROOT / ".env")
    analysis = json.loads(args.analysis_file.read_text(encoding="utf-8"))
    if analysis.get("next_action") != "draft_quote_for_manual_review":
        raise SystemExit("Analysis is not ready for a quote document.")

    args.out_dir.mkdir(parents=True, exist_ok=True)
    safe_name = "".join(ch if ch.isalnum() or ch in ("-", "_") else "_" for ch in args.recipient_name).strip("_")
    docx_path = args.out_dir / f"{safe_name or 'customer'}_solar_quote.docx"
    result_path = args.out_dir / f"{safe_name or 'customer'}_signwell_result.json"
    build_polished_quote_docx(docx_path, analysis, args.recipient_name)

    payload = build_signwell_payload(
        f"Solar quote - {args.recipient_name}",
        docx_path,
        args.recipient_email,
        args.recipient_name,
        draft=not args.send,
        test_mode=args.test_mode,
        embedded_signing=args.embedded,
    )
    result: dict[str, Any] = {
        "mode": "dry-run",
        "draft": not args.send,
        "test_mode": args.test_mode,
        "docx_path": str(docx_path),
        "document_name": payload["name"],
        "recipient_email": args.recipient_email,
    }
    if args.live_create:
        api_key = os.environ.get("SIGNWELL_API_KEY")
        if not api_key:
            raise SystemExit("SIGNWELL_API_KEY is missing. Add it to .env before --live-create.")
        response = create_document(api_key, os.environ.get("SIGNWELL_BASE_URL", DEFAULT_BASE_URL).rstrip("/"), payload)
        result.update({"mode": "live", "response": response, "document_id": response.get("id")})
        recipients = response.get("recipients") or []
        if recipients:
            result["signing_url"] = recipients[0].get("embedded_signing_url") or recipients[0].get("signing_url")
    result_path.write_text(json.dumps(result, indent=2), encoding="utf-8")
    print(json.dumps(result, indent=2))


if __name__ == "__main__":
    main()
